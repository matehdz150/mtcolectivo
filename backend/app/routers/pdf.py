# app/routers/pdf.py
from fastapi import APIRouter, UploadFile, File, Form, Depends, HTTPException
from fastapi.responses import StreamingResponse, PlainTextResponse
from sqlalchemy.orm import Session
from io import BytesIO
import os
import subprocess
import re
from typing import Any, Dict
from docxcompose.composer import Composer
from app.database import SessionLocal
from app.models import Order
from app.pdf_utils import docx_to_pdf_bytes
from app.deps import get_current_user   # rutas protegidas
from app.schemas import User            # (payload del usuario autenticado)
from docx import Document
from io import BytesIO
from app.pdf_utils import fill_docx_with_mapping
from app.pdf_utils import TEMPLATE_PATH
from docx.enum.section import WD_SECTION

router = APIRouter(prefix="/pdf", tags=["PDF"])

# ---- DB dependency ----
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# ---- Helpers ----
_num_re = re.compile(r"[^\d\-,.\s]")

def parse_num(val: Any) -> float:
    """
    Convierte a float tolerando '1,500', '1 500', '$1,500.00', etc.
    Vacíos o None => 0.0
    """
    if val is None:
        return 0.0
    s = str(val).strip()
    if not s:
        return 0.0
    s = _num_re.sub("", s)
    s = s.replace(" ", "")
    if "," in s and "." in s:
        s = s.replace(",", "")
    elif "," in s and "." not in s:
        s = s.replace(",", ".")
    try:
        return float(s)
    except ValueError:
        return 0.0

def generate_pdf_from_order(mapping: dict, texto_extra: str | None) -> bytes:

    # 1️⃣ Generar base DOCX
    base_docx = generate_docx_from_template(mapping)

    final_docx = base_docx

    # 2️⃣ Si hay texto extra → merge
    if texto_extra and texto_extra.strip():
        extra_docx = generate_extra_page(texto_extra)
        final_docx = merge_docx(base_docx, extra_docx)

    # 3️⃣ Convertir resultado final a PDF
    return docx_to_pdf_bytes(final_docx)

# ─────────────────────────────────────────────────────────────────────────────
# POST /pdf/from-data  → Genera PDF desde JSON (NO guarda en DB)
# Si te pasan subtotal/desc/abonado, recalculamos total y liquidar.
# ─────────────────────────────────────────────────────────────────────────────


@router.post("/from-data")
async def pdf_from_data(
    data: Dict[str, Any],
    current_user: User = Depends(get_current_user),
):
    try:
        def g(d: Dict[str, Any], k: str) -> str:
            for cand in (k, k.lower(), k.upper()):
                if cand in d:
                    v = d[cand]
                    return "" if v is None else str(v)
            return ""

        # 🔹 Números
        subtotal  = parse_num(g(data, "subtotal"))
        descuento = parse_num(g(data, "descuento"))
        abonado   = parse_num(g(data, "abonado"))

        total     = subtotal - descuento
        liquidar  = total - abonado

        # 🔹 Texto extra (🔥 aquí está la diferencia)
        texto_extra = g(data, "texto_extra")

        mapping = {
            "&NOMBRE&": g(data, "nombre"),
            "&FECHA&": g(data, "fecha"),
            "&DIR_SALIDA&": g(data, "dir_salida"),
            "&DIR_DESTINO&": g(data, "dir_destino"),
            "&HOR_IDA&": g(data, "hor_ida"),
            "&HOR_REGRESO&": g(data, "hor_regreso"),
            "&DURACION&": g(data, "duracion"),
            "&CAPACIDADU&": g(data, "capacidadu"),
            "&SUBTOTAL&": f"{subtotal:,.2f}",
            "&DESCUENTO&": f"{descuento:,.2f}",
            "&TOTAL&": f"{total:,.2f}",
            "&ABONADO&": f"{abonado:,.2f}",
            "&FECHA_ABONO&": g(data, "fecha_abono"),
            "&LIQUIDAR&": f"{liquidar:,.2f}",
        }

        # 🔥 USAMOS LA FUNCIÓN QUE HACE MERGE
        pdf_bytes = generate_pdf_from_order(
            mapping=mapping,
            texto_extra=texto_extra
        )

        return StreamingResponse(
            BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={"Content-Disposition": 'attachment; filename="orden.pdf"'},
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    

def _replace_tokens_in_paragraph(paragraph, mapping: dict) -> None:
    # Une runs -> reemplaza sobre el texto completo -> reescribe
    full = "".join(run.text for run in paragraph.runs)
    if not full:
        return

    changed = False
    for key, value in mapping.items():
        if key in full:
            full = full.replace(key, value)
            changed = True

    if not changed:
        return

    # Limpia todos los runs y deja el texto en el primero
    for run in paragraph.runs:
        run.text = ""
    paragraph.runs[0].text = full


def generate_docx_from_template(mapping: dict) -> bytes:
    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError("No se encontró la plantilla")

    with open(TEMPLATE_PATH, "rb") as f:
        tpl_bytes = f.read()

    return fill_docx_with_mapping(tpl_bytes, mapping)

EXTRA_TEMPLATE_PATH = os.path.join(
    os.path.dirname(__file__),
    "../../PlantillaExtra.docx"
)

def generate_extra_page(texto: str) -> bytes:
    doc = Document(EXTRA_TEMPLATE_PATH)

    # 🔥 FORZAR NUEVA PÁGINA AL INICIO
    doc.paragraphs[0].insert_paragraph_before().add_run().add_break()

    for paragraph in doc.paragraphs:
        if "&TEXTO_EXTRA&" in paragraph.text:
            paragraph.text = paragraph.text.replace("&TEXTO_EXTRA&", texto)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()

def merge_docx(base_bytes: bytes, extra_bytes: bytes) -> bytes:
    base_doc = Document(BytesIO(base_bytes))

    # 🔥 FORZAR NUEVA SECCIÓN EN NUEVA PÁGINA
    base_doc.add_section(WD_SECTION.NEW_PAGE)

    composer = Composer(base_doc)

    extra_doc = Document(BytesIO(extra_bytes))
    composer.append(extra_doc)

    buffer = BytesIO()
    composer.save(buffer)
    buffer.seek(0)
    return buffer.read()

    
@router.post("/from-data-word")
async def word_from_data(
    data: Dict[str, Any],
    current_user: User = Depends(get_current_user),
):
    try:
        def g(d: Dict[str, Any], k: str) -> str:
            for cand in (k, k.lower(), k.upper()):
                if cand in d:
                    v = d[cand]
                    return "" if v is None else str(v)
            return ""

        # ========= Leer valores =========
        s_subtotal  = g(data, "subtotal")
        s_descuento = g(data, "descuento")
        s_abonado   = g(data, "abonado")

        subtotal  = parse_num(s_subtotal)
        descuento = parse_num(s_descuento)
        abonado   = parse_num(s_abonado)

        total    = subtotal - descuento
        liquidar = total - abonado

        mapping = {
            "&NOMBRE&": g(data, "nombre"),
            "&FECHA&": g(data, "fecha"),
            "&DIR_SALIDA&": g(data, "dir_salida"),
            "&DIR_DESTINO&": g(data, "dir_destino"),
            "&HOR_IDA&": g(data, "hor_ida"),
            "&HOR_REGRESO&": g(data, "hor_regreso"),
            "&DURACION&": g(data, "duracion"),
            "&CAPACIDADU&": g(data, "capacidadu"),
            "&SUBTOTAL&": f"{subtotal:.2f}",
            "&DESCUENTO&": f"{descuento:.2f}",
            "&TOTAL&": f"{total:.2f}",
            "&ABONADO&": f"{abonado:.2f}",
            "&FECHA_ABONO&": g(data, "fecha_abono"),
            "&LIQUIDAR&": f"{liquidar:.2f}",
        }

        docx_bytes = generate_docx_from_template(mapping)

        return StreamingResponse(
            BytesIO(docx_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": 'attachment; filename="orden.docx"'
            },
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))